from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "output" / "doc" / "lacquertutor_report_first_person.md"
OUTPUT_DOCX = ROOT / "output" / "doc" / "lacquertutor_report_first_person.docx"

HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[(?P<alt>.*?)]\((?P<path>.*?)\)$")


def ensure_dirs() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    for style_name, size in (("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)

    if "Caption" not in doc.styles:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Microsoft YaHei"
        caption.font.size = Pt(9.5)
    else:
        caption = doc.styles["Caption"]
        caption.font.name = "Microsoft YaHei"
        caption.font.size = Pt(9.5)


def add_image(doc: Document, image_path: Path, alt_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(16.2))
    if alt_text:
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = 1
        caption.add_run(alt_text)


def build_docx() -> None:
    ensure_dirs()
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    list_buffer: list[str] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for item in list_buffer:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(item)
        list_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_list()
            continue

        heading = HEADING_RE.match(line)
        if heading:
            flush_list()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                if doc.paragraphs:
                    doc.add_section(WD_SECTION_START.CONTINUOUS)
                doc.add_heading(text, level=0)
            else:
                doc.add_heading(text, level=level - 1)
            continue

        image = IMAGE_RE.match(line.strip())
        if image:
            flush_list()
            image_path = (SOURCE_MD.parent / image.group("path")).resolve()
            add_image(doc, image_path, image.group("alt").strip())
            continue

        if line.startswith("- "):
            list_buffer.append(line[2:].strip())
            continue

        flush_list()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(line)

    flush_list()
    doc.save(OUTPUT_DOCX)
    print(f"Saved DOCX to {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx()
